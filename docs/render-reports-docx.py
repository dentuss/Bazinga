#!/usr/bin/env python3
"""
Render bachelor's-report Markdown files to Google-Docs-friendly DOCX.

Design constraints from the latest user feedback:
- single serif typeface throughout the body (Liberation Serif), 14 pt;
- highlighted runs (bold/italic) and monospaced code keep their styling;
- code snippets render as plain monospaced paragraphs with no border /
  shading / picture-like frame, divided by "Приклад N — …" captions
  (the captions are already part of the markdown);
- comparison tables are bordered and centered;
- title page mirrors the reference report layout.

The script parses the same markdown the PDF renderer uses, then walks the
parsed structure to emit Word XML via python-docx. Markdown features used
here are intentionally narrow (headings, paragraphs, bold/italic/inline
code, bullet/numbered lists, GFM tables, fenced code blocks); anything
more exotic is rendered as plain text.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent

SERIF = "Liberation Serif"      # one serif typeface for the entire body
MONO = "Liberation Mono"        # code only — still kept inside this single deck
BODY_PT = 14
CAPTION_PT = 13
CODE_PT = 11
H1_PT = 16
H2_PT = 15
H3_PT = 14


# ---------------------------------------------------------------------------
# Small helpers for tweaking python-docx XML where the public API stops short
# ---------------------------------------------------------------------------

def _set_cell_border(cell, color="000000", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_run_font(run, *, name=SERIF, size_pt=BODY_PT, bold=None, italic=None, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    # East-Asian font hint keeps Cyrillic mapped to the same family in MS Word
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _new_paragraph(doc, *, align=None, indent_first=False, space_after=2,
                   line_rule_auto=True, line_spacing=1.15, keep_with_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if line_rule_auto:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if indent_first:
        pf.first_line_indent = Cm(1.25)
    if align is not None:
        p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    return p


# ---------------------------------------------------------------------------
# Inline markdown → runs (bold / italic / `code`) on a target paragraph
# ---------------------------------------------------------------------------

# A simple, robust inline parser. Markdown bold uses **…** (we also accept __…__);
# italic uses *…* or _…_; inline code uses `…`. Sequences are non-overlapping
# and parsed greedily left-to-right.
_INLINE = re.compile(
    r"(\*\*[^*\n]+?\*\*|__[^_\n]+?__|\*[^*\n]+?\*|_[^_\n]+?_|`[^`\n]+?`)"
)


def _add_inline(paragraph, text: str, *, base_size=BODY_PT):
    """Walk the inline markdown of `text`, append styled runs to `paragraph`."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], size_pt=base_size)
        token = m.group(0)
        if token.startswith("**") or token.startswith("__"):
            _add_run(paragraph, token[2:-2], bold=True, size_pt=base_size)
        elif token.startswith("`"):
            _add_run(paragraph, token[1:-1], mono=True, size_pt=base_size)
        else:
            _add_run(paragraph, token[1:-1], italic=True, size_pt=base_size)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], size_pt=base_size)


def _add_run(paragraph, text: str, *, bold=False, italic=False, mono=False,
             size_pt=BODY_PT, color=None):
    if not text:
        return
    run = paragraph.add_run(text)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    _set_run_font(
        run,
        name=MONO if mono else SERIF,
        size_pt=size_pt,
        bold=bold or None,
        italic=italic or None,
    )


# ---------------------------------------------------------------------------
# Markdown parsing: we tokenise into a small set of block types
# ---------------------------------------------------------------------------

TABLE_SEP = re.compile(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def parse_blocks(md: str):
    """Yield (type, payload) tuples for top-level blocks in `md`."""
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Front-matter / horizontal rule
        if line.strip() == "---":
            yield ("hr", None)
            i += 1
            continue

        # Headings — H1 / H2 / H3
        if line.startswith("# "):
            yield ("h1", line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            yield ("h2", line[3:].strip())
            i += 1
            continue
        if line.startswith("### "):
            yield ("h3", line[4:].strip())
            i += 1
            continue

        # Fenced code block ```lang … ```
        if line.startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].startswith("```"):
                buf.append(lines[j])
                j += 1
            yield ("code", "\n".join(buf))
            i = j + 1
            continue

        # Tables — header row | --- | --- |
        if "|" in line and i + 1 < len(lines) and TABLE_SEP.match(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            j = i + 2
            rows = []
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            yield ("table", (header, rows))
            i = j
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            j = i
            items = []
            while j < len(lines) and re.match(r"^\s*[-*]\s+", lines[j]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[j]))
                j += 1
            yield ("ul", items)
            i = j
            continue

        # Numbered list
        if re.match(r"^\s*\d+[.)]\s+", line):
            j = i
            items = []
            while j < len(lines) and re.match(r"^\s*\d+[.)]\s+", lines[j]):
                items.append(re.sub(r"^\s*\d+[.)]\s+", "", lines[j]))
                j += 1
            yield ("ol", items)
            i = j
            continue

        # Paragraph: collect consecutive non-blank lines that are not block starters
        if line.strip():
            j = i
            chunk = []
            while j < len(lines) and lines[j].strip() \
                    and not lines[j].startswith(("# ", "## ", "### ", "```")) \
                    and not re.match(r"^\s*[-*]\s+", lines[j]) \
                    and not re.match(r"^\s*\d+[.)]\s+", lines[j]) \
                    and not ("|" in lines[j] and j + 1 < len(lines) and TABLE_SEP.match(lines[j + 1])) \
                    and lines[j].strip() != "---":
                chunk.append(lines[j])
                j += 1
            yield ("p", " ".join(chunk).strip())
            i = j
            continue

        # Blank line: skip
        i += 1


# ---------------------------------------------------------------------------
# Front matter (everything before the first H2) → custom title page
# ---------------------------------------------------------------------------

def render_title_page(doc, blocks, raw_md=""):
    """Find the title-page fields in the front matter and lay them out the
    way the reference report does. We look at the raw markdown (line by line)
    for the fixed-format header lines and only fall back to the parsed blocks
    for the team table."""
    fm = {}
    project_title = None
    team_rows = []
    author = supervisor = city_year = None

    # Raw line-by-line scan up to the first "---" or "## АНОТАЦІЯ" — that is
    # the title-page region of the markdown.
    for raw in raw_md.splitlines():
        s = raw.strip()
        if s == "---" or s.startswith("## АНОТАЦІЯ"):
            break
        plain = s.replace("**", "").strip()
        if plain == "Приватний заклад вищої освіти":
            fm["inst_top"] = plain
        elif plain == "Одеський технологічний університет «ШАГ»":
            fm["inst_main"] = plain
        elif plain.startswith("Кафедра"):
            fm["department"] = plain
        elif plain.startswith("на здобуття"):
            fm["degree"] = plain
        elif plain.startswith("зі спеціальності"):
            fm["specialty"] = plain
        elif s.startswith("## "):
            project_title = s[3:].strip()
        elif s.startswith("**Автор звіту:**"):
            author = s.split("**Автор звіту:**", 1)[1].strip()
        elif s.startswith("**Керівник:**"):
            supervisor = s.split("**Керівник:**", 1)[1].strip()
        elif re.match(r"^Одеса\s*[–-]\s*\d{4}$", s):
            city_year = s

    for kind, payload in blocks:
        if kind == "h2" and payload.strip().upper() == "АНОТАЦІЯ":
            break  # front matter ends here
        if kind == "table":
            header, rows = payload
            if header[:2] == ["№", "П.І.Б."]:
                team_rows = rows

    # Header (centered)
    for line in (fm.get("inst_top"), fm.get("inst_main"), fm.get("department")):
        if not line:
            continue
        p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        _add_run(p, line, bold=(line == fm.get("inst_main")), size_pt=BODY_PT)

    # Spacer
    for _ in range(4):
        _new_paragraph(doc, space_after=2)

    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_run(p, "випускна кваліфікаційна робота бакалавра", size_pt=BODY_PT)

    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_run(p, project_title or "", bold=True, size_pt=BODY_PT)

    for line in (fm.get("degree"), fm.get("specialty")):
        if not line:
            continue
        p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        _add_run(p, line, size_pt=BODY_PT)

    # Spacer
    for _ in range(3):
        _new_paragraph(doc, space_after=2)

    # "Виконавці проєкту:" + team table
    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_run(p, "Виконавці проєкту:", size_pt=BODY_PT)

    if team_rows:
        table = doc.add_table(rows=1 + len(team_rows), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = (Cm(1.2), Cm(8.5), Cm(4.0), Cm(2.5))
        hdr = table.rows[0].cells
        for col, label in enumerate(("№", "П.І.Б.", "Ролі", "Група")):
            hdr[col].text = ""
            cp = hdr[col].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(cp, label, bold=True, size_pt=BODY_PT)
            _set_cell_border(hdr[col])
            hdr[col].width = widths[col]
        for r, row in enumerate(team_rows, start=1):
            cells = table.rows[r].cells
            for c, val in enumerate(row[:4]):
                cells[c].text = ""
                cp = cells[c].paragraphs[0]
                cp.alignment = (WD_ALIGN_PARAGRAPH.CENTER if c in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT)
                _add_run(cp, val, size_pt=BODY_PT)
                _set_cell_border(cells[c])
                cells[c].width = widths[c]

    # Author
    if author:
        _new_paragraph(doc, space_after=2)
        p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        _add_run(p, f"Автор звіту: {author}", bold=True, size_pt=BODY_PT)

    # Supervisor mini-table
    if supervisor:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = (Cm(4.5), Cm(11.0))
        for col, (label, val) in enumerate((("Керівник", supervisor),)):
            pass
        c1, c2 = table.rows[0].cells
        c1.text = ""
        cp = c1.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(cp, "Керівник", size_pt=BODY_PT)
        _set_cell_border(c1)
        c1.width = widths[0]
        c2.text = ""
        cp = c2.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(cp, supervisor, size_pt=BODY_PT)
        _set_cell_border(c2)
        c2.width = widths[1]

    # Spacer + city/year
    for _ in range(3):
        _new_paragraph(doc, space_after=2)
    if city_year:
        p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        _add_run(p, city_year, size_pt=BODY_PT)
    # No manual page break here: the first body heading (АНОТАЦІЯ) carries
    # page-break-before, which keeps the title on its own (unnumbered) page.


# ---------------------------------------------------------------------------
# Body rendering
# ---------------------------------------------------------------------------

def _styled_heading(doc, text, *, level, align, size_pt, page_break, upper, bookmark=None):
    """Add a heading using the built-in Word Heading style (so the ЗМІСТ TOC
    field can pick it up), then override look to single-serif black."""
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(6 if level == 1 else 4)
    pf.space_after = Pt(6 if level == 1 else 4)
    pf.line_spacing = 1.15
    pf.keep_with_next = True
    pf.page_break_before = page_break
    p.alignment = align
    if bookmark:
        _add_bookmark(p, bookmark)
    _add_run(p, text.upper() if upper else text, bold=True, size_pt=size_pt, color="000000")
    return p


def render_body(doc, blocks):
    # Pre-compute the static ЗМІСТ once, using estimated page numbers from a
    # height-accounting pass over all body blocks. This makes the TOC plain
    # text (no live Word/Google Docs fields), exactly as requested.
    toc_entries = _collect_toc_entries(blocks)
    h2_seen = False
    skip_next_ul = False
    for kind, payload in blocks:
        # Skip the whole front matter (project title, team table, author,
        # supervisor) — it is laid out by render_title_page. The body proper
        # begins at the "АНОТАЦІЯ" heading.
        if not h2_seen:
            if kind == "h2" and payload.strip().upper() == "АНОТАЦІЯ":
                h2_seen = True
            else:
                continue

        if kind == "hr":
            continue  # the markdown horizontal rule is a structural separator only
        if kind == "h1":
            _styled_heading(doc, payload, level=1, align=WD_ALIGN_PARAGRAPH.CENTER,
                            size_pt=H1_PT, page_break=True, upper=True)
        elif kind == "h2":
            bookmark = None
            up = payload.strip().upper()
            # The ЗМІСТ section becomes a live table-of-contents field.
            if up == "ЗМІСТ":
                _styled_heading(doc, payload, level=1, align=WD_ALIGN_PARAGRAPH.CENTER,
                                size_pt=H1_PT, page_break=True, upper=True)
                _render_static_toc(doc, toc_entries)
                skip_next_ul = True
                continue
            if up.startswith("ДОДАТОК А"):
                bookmark = "dodatok_a"
            elif up.startswith("ДОДАТОК Б"):
                bookmark = "dodatok_b"
            _styled_heading(doc, payload, level=1, align=WD_ALIGN_PARAGRAPH.CENTER,
                            size_pt=H1_PT, page_break=True, upper=True, bookmark=bookmark)
        elif kind == "h3":
            _styled_heading(doc, payload, level=2, align=WD_ALIGN_PARAGRAPH.LEFT,
                            size_pt=H3_PT, page_break=False, upper=False)
        elif kind == "p":
            render_paragraph(doc, payload)
        elif kind == "ul":
            if skip_next_ul:
                skip_next_ul = False
                continue
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                _add_inline(p, item)
        elif kind == "ol":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                _add_inline(p, item)
        elif kind == "table":
            render_table(doc, *payload)
        elif kind == "code":
            render_code_block(doc, payload)


# --- Static ЗМІСТ -----------------------------------------------------------
# We render the TOC as plain paragraphs with a dot-leader tab and a numeric
# page number instead of a Word TOC field. The numbers are estimated by a
# simple height accounting over the parsed markdown blocks — it matches the
# layout closely enough for a print-ready document, and the user can adjust
# any line by hand without any field updates.

# Body geometry assumptions (used only for the page-number estimator):
#   • A4, margins top/bottom 2.0 cm, left 2.5 cm, right 1.5 cm
#   • text area ≈ 25.7 cm × 17.0 cm
#   • body 14 pt at 1.15 line spacing  → ~16.1 pt per line → ≈ 45 lines
#   • code 11 pt at 1.15 line spacing  → ~12.7 pt per line
# We use a slightly more conservative cap to allow for paragraph-after gaps,
# captions and headings that don't quite fit the average.
_PAGE_LINES = 36          # usable body lines per page after spacing overhead
_CHARS_PER_LINE = 68      # approx. for 14 pt over 17 cm column
_CODE_LINE_RATIO = 0.85   # 11 pt over 14 pt with similar leading


def _collect_toc_entries(blocks):
    """Walk parsed blocks and assign each TOC-able heading an estimated page.

    Returns a list of tuples (level, text, page_no), where:
      • level 1 → top-level section (АНОТАЦІЯ, ЗМІСТ, ВСТУП, РОЗДІЛ N, …),
      • level 2 → numbered subsection (1.1, 2.1, …).
    """
    entries: list[tuple[int, str, int]] = []
    started = False                    # body proper starts at АНОТАЦІЯ
    in_appendix = False                # suppress subsection entries inside an appendix
    page = 1                            # title page is page 1
    lines = 0                           # lines accumulated on the current page

    def page_break():
        nonlocal page, lines
        page += 1
        lines = 0

    def consume_lines(n: int):
        """Account for `n` body-equivalent lines, breaking pages as needed."""
        nonlocal page, lines
        if n <= 0:
            return
        while n > 0:
            free = _PAGE_LINES - lines
            if n <= free:
                lines += n
                return
            n -= free
            page_break()

    for kind, payload in blocks:
        if not started:
            if kind == "h2" and payload.strip().upper() == "АНОТАЦІЯ":
                page_break()                # АНОТАЦІЯ lands on page 2
                started = True
                entries.append((1, payload.strip(), page))
                lines = 3                   # heading itself takes ~3 lines
            continue

        if kind == "h2":
            page_break()
            text = payload.strip()
            entries.append((1, text, page))
            in_appendix = text.upper().startswith("ДОДАТОК")
            lines = 3
        elif kind == "h3":
            # Subsections don't force a page break; if we're near the bottom,
            # the heading will spill onto the next page anyway.
            if lines + 4 > _PAGE_LINES:
                page_break()
            # Match the reference: subsections inside the appendix are not
            # listed in the ЗМІСТ — only the appendix heading itself appears.
            if not in_appendix:
                entries.append((2, payload.strip(), page))
            lines += 3
        elif kind == "p":
            words = max(1, len(payload.split()))
            est_lines = max(1, (words * 7 + _CHARS_PER_LINE - 1) // _CHARS_PER_LINE) + 1
            consume_lines(est_lines)
        elif kind == "ul" or kind == "ol":
            for item in payload:
                words = max(1, len(item.split()))
                est_lines = max(1, (words * 7 + _CHARS_PER_LINE - 1) // _CHARS_PER_LINE)
                consume_lines(est_lines)
        elif kind == "code":
            code_lines = payload.count("\n") + 2
            consume_lines(int(code_lines * _CODE_LINE_RATIO) + 1)
        elif kind == "table":
            header, rows = payload
            # ≈ 2 body-line equivalents per table row + small gap.
            consume_lines(2 * (len(rows) + 1) + 1)

    return entries


def _render_static_toc(doc, entries):
    """Render the ЗМІСТ as plain paragraphs: text — dot leader — page number.

    Each paragraph carries a right-aligned tab stop with a dot leader at the
    right edge of the text column, so titles and numbers align like in the
    reference report and a typical autoreferat.
    """
    for level, text, page in entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.15
        if level == 2:
            pf.left_indent = Cm(0.8)
        pf.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        _add_run(p, text, bold=(level == 1), size_pt=BODY_PT)

        # Tab character — the tab stop above pulls the page number to the
        # right edge and fills the gap with dots.
        tab_run = p.add_run()
        _set_run_font(tab_run, size_pt=BODY_PT)
        tab_elem = OxmlElement("w:tab")
        tab_run._r.append(tab_elem)

        _add_run(p, str(page), bold=(level == 1), size_pt=BODY_PT)


def render_paragraph(doc, text: str):
    """Body paragraphs. Recognise our caption convention so 'Приклад N — …'
    and 'Таблиця X — …' come out centred + italic + kept-with-next; turn
    references to appendices into internal hyperlinks."""
    is_caption = bool(re.match(r"^(Приклад|Таблиця|Рисунок|Лістинг)\s\S", text))
    if is_caption:
        p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
                           keep_with_next=True)
        _add_inline_italic(p, text, size_pt=CAPTION_PT)
        return
    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       indent_first=True, space_after=4)
    # Split on appendix references and render those as internal hyperlinks.
    pos = 0
    for m in _DODATOK_REF.finditer(text):
        anchor = _ANCHOR.get(m.group(1))
        if not anchor:
            continue
        if m.start() > pos:
            _add_inline(p, text[pos:m.start()])
        _add_internal_hyperlink(p, m.group(0), anchor)
        pos = m.end()
    if pos < len(text):
        _add_inline(p, text[pos:])


def _add_inline_italic(paragraph, text: str, *, size_pt=CAPTION_PT):
    """Same as _add_inline but italicises plain text runs (captions)."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], italic=True, size_pt=size_pt)
        token = m.group(0)
        if token.startswith("**") or token.startswith("__"):
            _add_run(paragraph, token[2:-2], bold=True, italic=True, size_pt=size_pt)
        elif token.startswith("`"):
            _add_run(paragraph, token[1:-1], mono=True, italic=True, size_pt=size_pt)
        else:
            _add_run(paragraph, token[1:-1], italic=True, size_pt=size_pt)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], italic=True, size_pt=size_pt)


def render_table(doc, header, rows):
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for c, label in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cp, label, bold=True, size_pt=BODY_PT)
        _set_cell_border(cell)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row[:cols]):
            cell = table.rows[r].cells[c]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            # First column tends to be a label; centre the +/− cells.
            cp.alignment = (WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            _add_inline(cp, val, base_size=BODY_PT - 1)
            _set_cell_border(cell)
    # Tiny gap after the table.
    _new_paragraph(doc, space_after=4)


def render_code_block(doc, code: str):
    """Render fenced code as plain monospaced paragraphs — no border, no
    shading, no background. One paragraph per source line so Google Docs can
    flow it across pages cleanly."""
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15
        pf.left_indent = Cm(1.0)
        # Disable any default style background/borders inherited from the doc.
        run = p.add_run(line if line else " ")
        _set_run_font(run, name=MONO, size_pt=CODE_PT)
    _new_paragraph(doc, space_after=4)


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def _apply_default_style(doc):
    """Set Normal style to single-serif 14 pt so any stray paragraph we don't
    explicitly style still picks up the right typeface."""
    style = doc.styles["Normal"]
    style.font.name = SERIF
    style.font.size = Pt(BODY_PT)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), SERIF)


def _apply_page_setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)


def _add_field(paragraph, instr: str, placeholder: str = "", *, name=SERIF, size_pt=BODY_PT):
    """Insert a Word complex field (e.g. PAGE, TOC) into `paragraph`."""
    r1 = paragraph.add_run(); _set_run_font(r1, name=name, size_pt=size_pt)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._element.append(f1)
    r2 = paragraph.add_run(); _set_run_font(r2, name=name, size_pt=size_pt)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._element.append(it)
    r3 = paragraph.add_run(); _set_run_font(r3, name=name, size_pt=size_pt)
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "separate"); r3._element.append(f3)
    r4 = paragraph.add_run(placeholder); _set_run_font(r4, name=name, size_pt=size_pt)
    r5 = paragraph.add_run(); _set_run_font(r5, name=name, size_pt=size_pt)
    f5 = OxmlElement("w:fldChar"); f5.set(qn("w:fldCharType"), "end"); r5._element.append(f5)


def _add_page_number_footer(doc):
    """Bottom-right page number on every page except the title page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # default (non-first) footer — carries the page number
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_field(p, " PAGE ", "2", name=SERIF, size_pt=12)
    # first-page footer stays empty so the title page has no number
    fp = section.first_page_footer
    fp.is_linked_to_previous = False
    if not fp.paragraphs:
        fp.add_paragraph()


def _add_bookmark(paragraph, name: str):
    """Wrap a bookmark around the start of `paragraph` (for internal links)."""
    bid = str(abs(hash(name)) % 100000)
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), bid); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), bid)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph, text: str, anchor: str, *, size_pt=BODY_PT):
    """Append an internal hyperlink run (blue, underlined) pointing to a bookmark."""
    hl = OxmlElement("w:hyperlink"); hl.set(qn("w:anchor"), anchor)
    run = paragraph.add_run(text)
    _set_run_font(run, name=SERIF, size_pt=size_pt, color="1A4A8A")
    run.font.underline = True
    hl.append(run._element)
    paragraph._p.append(hl)


# Detects references to appendices, e.g. "Додаток А", "Додатку А", "Додатка Б".
_DODATOK_REF = re.compile(r"Додат(?:ок|ку|ка|кax)?\w*\s+([АБВГ])")
_ANCHOR = {"А": "dodatok_a", "Б": "dodatok_b", "В": "dodatok_v", "Г": "dodatok_g"}


def render(md_path: Path, docx_path: Path):
    md_text = md_path.read_text(encoding="utf-8")
    blocks = list(parse_blocks(md_text))
    doc = Document()
    _apply_default_style(doc)
    _apply_page_setup(doc)
    _add_page_number_footer(doc)
    render_title_page(doc, blocks, raw_md=md_text)
    render_body(doc, blocks)
    doc.save(docx_path)
    print(f"wrote {docx_path.name} ({docx_path.stat().st_size // 1024} KB)")


def main(argv):
    sources = [Path(p) for p in argv] if argv else sorted(ROOT.glob("bachelor-report-*.md"))
    for md_path in sources:
        render(md_path, md_path.with_suffix(".docx"))


if __name__ == "__main__":
    main(sys.argv[1:])
